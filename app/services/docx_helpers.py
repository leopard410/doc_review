import re

from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph

from app.models import ReviewCategory


CATEGORY_COLORS: dict[ReviewCategory, WD_COLOR_INDEX] = {
    ReviewCategory.SPELLING: WD_COLOR_INDEX.YELLOW,
    ReviewCategory.UNUSUAL: WD_COLOR_INDEX.TURQUOISE,
    ReviewCategory.TYPOGRAPHY: WD_COLOR_INDEX.BRIGHT_GREEN,
}

EMPHASIS_COLOR = WD_COLOR_INDEX.PINK
BLOCKQUOTE_COLOR = WD_COLOR_INDEX.GRAY_25


def _normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _find_phrase_span(paragraph_text: str, phrase: str) -> tuple[int, int] | None:
    if not phrase or not phrase.strip():
        return None

    direct = paragraph_text.find(phrase)
    if direct != -1:
        return direct, direct + len(phrase)

    normalized_para = _normalize_whitespace(paragraph_text)
    normalized_phrase = _normalize_whitespace(phrase)
    if not normalized_phrase:
        return None

    if normalized_phrase not in normalized_para:
        # Case-insensitive fallback
        lower_para = normalized_para.lower()
        lower_phrase = normalized_phrase.lower()
        idx = lower_para.find(lower_phrase)
        if idx == -1:
            return None
        normalized_phrase = normalized_para[idx : idx + len(lower_phrase)]

    pattern = re.escape(normalized_phrase).replace(r"\ ", r"\s+")
    match = re.search(pattern, paragraph_text, flags=re.IGNORECASE)
    if match:
        return match.start(), match.end()

    return None


def _build_char_map(paragraph: Paragraph) -> list[tuple[int, int]]:
    mapping: list[tuple[int, int]] = []
    for run_idx, run in enumerate(paragraph.runs):
        for char_idx, _ in enumerate(run.text):
            mapping.append((run_idx, char_idx))
    return mapping


def _highlight_via_runs(paragraph: Paragraph, start: int, end: int, color: WD_COLOR_INDEX) -> bool:
    char_map = _build_char_map(paragraph)
    if not char_map or end > len(char_map):
        return False

    run_indices: set[int] = set()
    for position in range(start, end):
        run_indices.add(char_map[position][0])

    for run_idx in run_indices:
        paragraph.runs[run_idx].font.highlight_color = color
    return True


def _clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)


def _append_hint(paragraph: Paragraph, hint: str) -> None:
    run = paragraph.add_run(f" [{hint}]")
    run.font.italic = True


def _rebuild_with_highlight(
    paragraph: Paragraph,
    start: int,
    end: int,
    color: WD_COLOR_INDEX,
    hint: str | None = None,
) -> None:
    text = paragraph.text
    before = text[:start]
    middle = text[start:end]
    after = text[end:]

    _clear_runs(paragraph)

    if before:
        paragraph.add_run(before)
    highlighted = paragraph.add_run(middle)
    highlighted.font.highlight_color = color
    if hint:
        _append_hint(paragraph, hint)
    if after:
        paragraph.add_run(after)


def highlight_phrase_in_paragraph(
    paragraph: Paragraph,
    phrase: str,
    color: WD_COLOR_INDEX,
    hint: str | None = None,
) -> bool:
    paragraph_text = paragraph.text
    if not paragraph_text.strip():
        return False

    span = _find_phrase_span(paragraph_text, phrase)
    if span is None:
        return False

    start, end = span

    if paragraph.runs and _highlight_via_runs(paragraph, start, end, color):
        if hint:
            _append_hint(paragraph, hint)
        return True

    _rebuild_with_highlight(paragraph, start, end, color, hint)
    return True


def highlight_in_range(
    paragraphs: list[Paragraph],
    start_idx: int,
    end_idx: int,
    phrase: str,
    color: WD_COLOR_INDEX,
    hint: str | None = None,
) -> bool:
    for idx in range(start_idx, end_idx):
        if highlight_phrase_in_paragraph(paragraphs[idx], phrase, color, hint):
            return True
    return False
