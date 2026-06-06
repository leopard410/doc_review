from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from app.models import ChapterAnalysis
from app.services.chapter_parser import Chapter, get_chapter_text
from app.services.docx_helpers import (
    BLOCKQUOTE_COLOR,
    CATEGORY_COLORS,
    EMPHASIS_COLOR,
    highlight_in_range,
)


def _find_insertion_paragraph(
    paragraphs: list[Paragraph],
    start_idx: int,
    end_idx: int,
    anchor: str | None,
) -> int:
    """Return paragraph index after which the closing heading should be inserted."""
    if end_idx <= start_idx:
        return max(0, start_idx)

    if anchor:
        for idx in range(start_idx, end_idx):
            if anchor in paragraphs[idx].text:
                return idx

    # Default: before the last non-empty paragraph in the chapter.
    for idx in range(end_idx - 1, start_idx - 1, -1):
        if paragraphs[idx].text.strip():
            return idx

    return max(0, end_idx - 1)


def _insert_closing_heading(
    document: Document,
    paragraphs: list[Paragraph],
    after_idx: int,
    heading_text: str,
) -> None:
    if after_idx < 0 or after_idx >= len(paragraphs):
        return

    ref = paragraphs[after_idx]
    new_p = OxmlElement("w:p")
    ref._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, ref._parent)

    for style_name in ("Heading 2", "Heading2", "heading 2"):
        try:
            new_paragraph.style = style_name
            break
        except KeyError:
            continue

    new_paragraph.add_run(heading_text)


def apply_chapter_annotations(
    document: Document,
    chapter: Chapter,
    analysis: ChapterAnalysis,
    closing_heading: str,
) -> dict[str, int | list[str]]:
    paragraphs = list(document.paragraphs)
    stats: dict[str, int | list[str]] = {
        "highlights_applied": 0,
        "highlights_missed": [],
        "closing_heading_inserted": 0,
    }
    missed: list[str] = []

    start = chapter.start_para_idx
    end = chapter.end_para_idx

    for item in analysis.review_items:
        color = CATEGORY_COLORS[item.category]
        label = f"{item.category.value}: {item.note}" if item.note else item.category.value
        if highlight_in_range(paragraphs, start, end, item.text, color, label):
            stats["highlights_applied"] = int(stats["highlights_applied"]) + 1
        else:
            missed.append(item.text)

    if chapter.is_odd:
        if analysis.emphasis_sentence:
            if highlight_in_range(
                paragraphs,
                start,
                end,
                analysis.emphasis_sentence,
                EMPHASIS_COLOR,
                "suggested emphasis",
            ):
                stats["highlights_applied"] = int(stats["highlights_applied"]) + 1
            else:
                missed.append(analysis.emphasis_sentence)

        if analysis.blockquote_sentence:
            if highlight_in_range(
                paragraphs,
                start,
                end,
                analysis.blockquote_sentence,
                BLOCKQUOTE_COLOR,
                "suggested block quote",
            ):
                stats["highlights_applied"] = int(stats["highlights_applied"]) + 1
            else:
                missed.append(analysis.blockquote_sentence)

    insert_after = _find_insertion_paragraph(
        paragraphs,
        start,
        end,
        analysis.closing_anchor,
    )
    _insert_closing_heading(document, paragraphs, insert_after, closing_heading)
    stats["closing_heading_inserted"] = 1

    stats["highlights_missed"] = missed
    return stats


def process_document(
    input_path: str,
    output_path: str,
    analyses: list[tuple[Chapter, ChapterAnalysis]],
    closing_heading: str,
) -> dict:
    document = Document(input_path)
    report: dict = {"chapters": []}

    for chapter, analysis in analyses:
        chapter_text = get_chapter_text(document, chapter)
        stats = apply_chapter_annotations(
            document,
            chapter,
            analysis,
            closing_heading,
        )
        report["chapters"].append(
            {
                "index": chapter.index,
                "title": chapter.title,
                "text_length": len(chapter_text),
                "review_items_found": len(analysis.review_items),
                **stats,
            }
        )

    document.save(output_path)
    return report
