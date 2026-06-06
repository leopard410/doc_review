"""Generate a small sample .docx for local testing."""

from pathlib import Path

from docx import Document


def main() -> None:
    doc = Document()
    doc.add_heading("Chapter One: The Beginning", level=1)
    doc.add_paragraph(
        "The journey began on a misty morning. Something felt unusal about the air, "
        "as if the world had shifted overnight. She whispered, 'We cannot go back,' "
        "and the words hung between them like smoke."
    )
    doc.add_paragraph(
        "This was the moment everything changed. Teh road ahead would test every "
        "assumption they had made. A single misprint in the map led them astray."
    )

    doc.add_heading("Chapter Two: The Crossing", level=1)
    doc.add_paragraph(
        "They crossed the river at dawn. The water was cold and unforgiving, "
        "but they pressed on without complaint."
    )
    doc.add_paragraph(
        "By evening, the village lights appeared on the horizon. Relief washed "
        "over the party in quiet waves."
    )

    doc.add_heading("Chapter Three: The Reckoning", level=1)
    doc.add_paragraph(
        "Confrontation was inevitable. He stood at the center of the square and "
        "declared, 'Truth will outlast every lie.' The crowd fell silent."
    )
    doc.add_paragraph(
        "No one moved. The silence itself became a kind of answer — a punctuation "
        "mark at the end of a long sentence."
    )

    out = Path("sample/manuscript.docx")
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Created {out}")


if __name__ == "__main__":
    main()
